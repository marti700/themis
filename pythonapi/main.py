from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, RichText
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from concordancia import (
    agree,
    denominar,
    ha,
    marital_status_es,
    party_block,
    pick,
    quien,
    resolve_denomination,
    sigue,
)

app = FastAPI(title="Themis Document API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMPLATES_DIR = Path(__file__).parent / "docxtemplates"
STATIC_DIR = Path(__file__).parent / "static"


@app.get("/templates")
def list_templates():
    return [f.stem for f in sorted(TEMPLATES_DIR.glob("*.docx"))]
_SIG_SENTINEL = "__SIGNATURES__"
_CLAUSES_SENTINEL = "__CLAUSES__"
_HEIRS_SENTINEL = "__HEREDEROS__"

# Spanish ordinal words for auto-numbering notarial-act clauses (PRIMERO, SEGUNDO, …).
_ORDINALS = [
    "PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO", "SÉPTIMO",
    "OCTAVO", "NOVENO", "DÉCIMO", "DÉCIMO PRIMERO", "DÉCIMO SEGUNDO",
    "DÉCIMO TERCERO", "DÉCIMO CUARTO", "DÉCIMO QUINTO", "DÉCIMO SEXTO",
    "DÉCIMO SÉPTIMO", "DÉCIMO OCTAVO", "DÉCIMO NOVENO", "VIGÉSIMO",
]


def ordinal_word(index: int) -> str:
    """0-based clause index → Spanish ordinal label (falls back to 'N-ÉSIMO')."""
    return _ORDINALS[index] if index < len(_ORDINALS) else f"{index + 1}-ÉSIMO"


class CustomerInfo(BaseModel):
    first_name: str
    last_name: str
    id_number: str
    address: str | None = None
    marital_status: str | None = None
    occupation: str | None = None
    nationality: str | None = None
    gender: str | None = None


class SellContractRequest(BaseModel):
    sellers: list[CustomerInfo]
    buyers: list[CustomerInfo]
    seller_denomination: str = "EL VENDEDOR"
    buyer_denomination: str = "EL COMPRADOR"
    property_description: str
    property_justification: str
    price_words: str
    price_rd: str
    originals_count_words: str
    originals_count: str
    city: str
    province: str
    day_words: str
    day_number: str
    month: str
    year_words: str
    year_number: str
    signers_list: str = ""
    firm_name: str
    firm_service: str
    notary_title_name: str
    notary_phone: str
    notary_municipality: str
    notary_registration: str


class RentContractRequest(BaseModel):
    owners: list[CustomerInfo]
    tenants: list[CustomerInfo]
    owner_denomination: str = "EL PROPIETARIO"
    tenant_denomination: str = "EL INQUILINO"
    property_description: str
    property_use: str = "vivienda"
    price_words: str
    price_rd: str
    deposit_words: str
    deposit_rd: str
    duration_words: str
    duration_number: str
    start_date: str
    end_date: str
    originals_count_words: str
    originals_count: str
    city: str
    province: str
    day_words: str
    day_number: str
    month: str
    year_words: str
    year_number: str
    signers_list: str = ""
    firm_name: str
    firm_service: str
    notary_title_name: str
    notary_phone: str
    notary_municipality: str
    notary_registration: str


class NotaryActBase(BaseModel):
    """Shared fields for actos auténticos where the notary is a party from the opening
    sentence and instrumental witnesses (testigos) sign the act."""

    # Firm / notary letterhead
    firm_name: str
    firm_service: str
    notary_title_name: str
    notary_phone: str
    notary_municipality: str
    notary_registration: str
    # Notary as an auténtico party (opening "Por ante mí, …")
    notary_nationality: str = ""
    notary_marital_status: str = ""
    notary_cedula: str = ""
    notary_office: str = ""
    # Instrumental witnesses (customer-picked, N)
    witnesses: list[CustomerInfo] = []
    # Place / date / closing
    city: str
    province: str
    day_words: str
    day_number: str
    month: str
    year_words: str
    year_number: str
    originals_count_words: str = "DOS"
    originals_count: str = "2"
    signers_list: str = ""


class PoderEspecialRequest(NotaryActBase):
    poderdantes: list[CustomerInfo]
    poderdante_denomination: str = "EL PODERDANTE"
    # Apoderado — the person receiving the power (customer-picked, N).
    apoderados: list[CustomerInfo] = []
    apoderado_denomination: str = "EL APODERADO"
    power_object: str


class ClauseActRequest(NotaryActBase):
    """Shared shape for acts where N declarants testify and the body is a dynamic,
    auto-numbered list of clauses (Declaración Jurada, Acto de Notoriedad)."""

    declarants: list[CustomerInfo]
    declarant_denomination: str = "EL COMPARECIENTE"
    clauses: list[str] = []


class DeclaracionJuradaRequest(ClauseActRequest):
    act_number: str = ""
    time: str = ""


class ActoNotoriedadRequest(ClauseActRequest):
    pass


class HeirInfo(BaseModel):
    """One heir enumerated in a Determinación de Herederos — richer than
    CustomerInfo because the act must state each heir's filiation (birth data)."""

    first_name: str
    last_name: str
    id_number: str = ""
    nationality: str = ""
    marital_status: str = ""
    occupation: str = ""
    address: str = ""
    birth_date: str = ""
    birth_record: str = ""
    gender: str | None = None


class DeterminacionHerederosRequest(NotaryActBase):
    declarants: list[CustomerInfo]
    declarant_denomination: str = "LOS DECLARANTES"
    deceased_name: str
    deceased_nationality: str = ""
    deceased_marital_status: str = ""
    deceased_occupation: str = ""
    deceased_id_number: str = ""
    deceased_birth_date: str = ""
    deceased_death_date: str = ""
    deceased_death_record: str = ""
    deceased_origin_place: str = ""
    deceased_spouse_name: str = ""
    heirs: list[HeirInfo] = []
    clauses: list[str] = []


class AutorizacionViajeMenorRequest(NotaryActBase):
    authorizers: list[CustomerInfo]
    authorizer_denomination: str = "EL AUTORIZANTE"
    minor_name: str
    minor_birth_date: str = ""
    minor_birth_record: str = ""
    minor_passport: str = ""
    destination_country: str
    acceptors: list[CustomerInfo] = []
    acceptor_denomination: str = "EL ACEPTANTE"


def _set_table_borders_invisible(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "nil")
        tblBorders.append(border_el)


def inject_signatures(doc: Document, parties: list[tuple[str, str]]):
    """Replace the sentinel paragraph with two-column signature tables (name bold+underlined, role below)."""
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if _SIG_SENTINEL in paragraph.text:
            target_paragraph = paragraph
            paragraph.text = ""
            break

    if target_paragraph is None:
        return

    # Build pairs; insert in reverse so addnext() preserves top-to-bottom order.
    pairs = [parties[i : i + 2] for i in range(0, len(parties), 2)]
    for chunk in reversed(pairs):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _set_table_borders_invisible(table)
        target_paragraph._p.addnext(table._tbl)

        for col_index, (name, role) in enumerate(chunk):
            cell = table.cell(0, col_index)

            name_para = cell.paragraphs[0]
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.underline = True

            role_para = cell.add_paragraph()
            role_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            role_para.add_run(role)

        spacer = doc.add_paragraph()
        table._tbl.addnext(spacer._p)


def inject_clauses(doc: Document, clauses: list[str], start_index: int = 0):
    """Replace the __CLAUSES__ sentinel paragraph with one justified paragraph per
    clause, each opening with a bold auto-numbered ordinal (PRIMERO:, SEGUNDO:, …).
    `start_index` continues the numbering after any clauses the template hardcodes
    (e.g. Determinación de Herederos fixes PRIMERO/SEGUNDO and starts dynamic
    clauses at TERCERO, index 2)."""
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if _CLAUSES_SENTINEL in paragraph.text:
            target_paragraph = paragraph
            paragraph.text = ""
            break

    if target_paragraph is None:
        return

    items = [c.strip() for c in clauses if c and c.strip()]
    # Insert in reverse so addnext() preserves PRIMERO → último order.
    for local_index in reversed(range(len(items))):
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ordinal_run = para.add_run(f"{ordinal_word(start_index + local_index)}: ")
        ordinal_run.bold = True
        para.add_run(items[local_index])
        target_paragraph._p.addnext(para._p)


def inject_heirs(doc: Document, heirs: list[HeirInfo]):
    """Replace the __HEREDEROS__ sentinel paragraph with one justified, numbered
    ("1-", "2-", …) paragraph per heir, stating their generals and filiation."""
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if _HEIRS_SENTINEL in paragraph.text:
            target_paragraph = paragraph
            paragraph.text = ""
            break

    if target_paragraph is None:
        return

    for index in reversed(range(len(heirs))):
        h = heirs[index]
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        prefix_run = para.add_run(f"{index + 1}- ")
        prefix_run.bold = True
        name_run = para.add_run(f"{h.first_name} {h.last_name}".upper())
        name_run.bold = True
        details = ""
        if h.nationality:
            details += f", {h.nationality}"
        details += ", mayor de edad"
        ms = marital_status_es(h.marital_status, h.gender)
        if ms:
            details += f", {ms}"
        if h.occupation:
            details += f", {h.occupation}"
        if h.id_number:
            portador = agree("portador", "portadora", h.gender)
            details += f", {portador} de la cédula de identidad y electoral No. {h.id_number}"
        if h.address:
            domiciliado = agree("domiciliado", "domiciliada", h.gender)
            details += f", {domiciliado} y residente en {h.address}"
        if h.birth_date:
            nacido = agree("nacido", "nacida", h.gender)
            details += f", {nacido} en fecha {h.birth_date}"
        if h.birth_record:
            details += f", según {h.birth_record}"
        details += ";"
        para.add_run(details)
        target_paragraph._p.addnext(para._p)


def _bold_upper(text: str) -> RichText:
    """Party name for a `{{r VAR}}` placeholder: bold and upper-cased, as notarial
    acts render the parties' names in the body text."""
    return RichText(text.upper(), bold=True)


def _notary_context(req: NotaryActBase) -> dict:
    """Placeholders shared by every acto auténtico (firm/notary letterhead, the notary
    as opening party, witnesses, place/date, and the signature sentinel)."""
    return {
        "BUFETE_NOMBRE": req.firm_name,
        "BUFETE_SERVICIO": req.firm_service,
        "NOTARIO_TITULO_NOMBRE": req.notary_title_name,
        "NOTARIO_TITULO_NOMBRE_MAYUS": req.notary_title_name.upper(),
        "NOTARIO_TELEFONO": req.notary_phone,
        "NOTARIO_MUNICIPIO": req.notary_municipality,
        "NOTARIO_MATRICULA": req.notary_registration,
        "NOTARIO_NACIONALIDAD": req.notary_nationality,
        "NOTARIO_ESTADO_CIVIL": req.notary_marital_status,
        "NOTARIO_CEDULA": req.notary_cedula,
        "NOTARIO_ESTUDIO": req.notary_office,
        "TESTIGO_BLOQUE": party_block(req.witnesses),
        "CIUDAD_CONTRATO": req.city,
        "PROVINCIA_CONTRATO": req.province,
        "DIA_LETRAS": req.day_words,
        "DIA_NUM": req.day_number,
        "MES": req.month,
        "ANIO_LETRAS": req.year_words,
        "ANIO_NUM": req.year_number,
        "CANTIDAD_ORIGINALES_LETRAS": req.originals_count_words,
        "CANTIDAD_ORIGINALES": req.originals_count,
        "FIRMANTES_LISTA": req.signers_list,
        "SIGNATURES": _SIG_SENTINEL,
    }


def _clause_act_context(req: ClauseActRequest) -> dict:
    """Placeholders shared by Declaración Jurada and Acto de Notoriedad: the
    declarant/compareciente party block and the dynamic-clauses sentinel."""
    n = len(req.declarants)
    return {
        "COMPARECIENTE_BLOQUE": party_block(req.declarants),
        "COMPARECIENTE_HA": ha(n),
        "COMPARECIENTE_QUIEN": quien(n),
        "COMPARECIENTE_SIGUE": sigue(n),
        "COMPARECIENTE_SE_DENOMINARA": denominar(n),
        "DENOMINACION_COMPARECIENTE": resolve_denomination(
            req.declarant_denomination, "COMPARECIENTE", req.declarants
        ),
        "CLAUSULAS": _CLAUSES_SENTINEL,
    }


def _act_response(doc: Document, filename: str) -> StreamingResponse:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/documents/sell_contract")
def generate_sell_contract(req: SellContractRequest):
    template_path = TEMPLATES_DIR / "sell_contract.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    # Pass 1 — docxtpl renders all text placeholders.
    # SIGNATURES is set to a sentinel so we can locate the paragraph in pass 2.
    tpl = DocxTemplate(template_path)
    context = {
        "BUFETE_NOMBRE": req.firm_name,
        "BUFETE_SERVICIO": req.firm_service,
        "NOTARIO_TITULO_NOMBRE": req.notary_title_name,
        "NOTARIO_TITULO_NOMBRE_MAYUS": req.notary_title_name.upper(),
        "NOTARIO_TELEFONO": req.notary_phone,
        "NOTARIO_MUNICIPIO": req.notary_municipality,
        "NOTARIO_MATRICULA": req.notary_registration,
        "VENDEDOR_BLOQUE": party_block(req.sellers),
        "VENDEDOR_QUIEN": quien(len(req.sellers)),
        "VENDEDOR_SIGUE": sigue(len(req.sellers)),
        "VENDEDOR_SE_DENOMINARA": denominar(len(req.sellers)),
        "DENOMINACION_VENDEDORES": resolve_denomination(req.seller_denomination, "VENDEDOR", req.sellers),
        "COMPRADOR_BLOQUE": party_block(req.buyers),
        "COMPRADOR_QUIEN": quien(len(req.buyers)),
        "COMPRADOR_SIGUE": sigue(len(req.buyers)),
        "COMPRADOR_SE_DENOMINARA": denominar(len(req.buyers)),
        "DENOMINACION_COMPRADOR": resolve_denomination(req.buyer_denomination, "COMPRADOR", req.buyers),
        "DESCRIPCION_INMUEBLE": req.property_description,
        "JUSTIFICACION_PROPIEDAD": req.property_justification,
        "PRECIO_LETRAS": req.price_words,
        "PRECIO_RD": req.price_rd,
        "CANTIDAD_ORIGINALES_LETRAS": req.originals_count_words,
        "CANTIDAD_ORIGINALES": req.originals_count,
        "CIUDAD_CONTRATO": req.city,
        "PROVINCIA_CONTRATO": req.province,
        "DIA_LETRAS": req.day_words,
        "DIA_NUM": req.day_number,
        "MES": req.month,
        "ANIO_LETRAS": req.year_words,
        "ANIO_NUM": req.year_number,
        "SIGNATURES": _SIG_SENTINEL,
        "FIRMANTES_LISTA": req.signers_list,
    }
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    # Pass 2 — python-docx finds the sentinel and injects the signature table.
    doc = Document(intermediate)
    parties = [
        (f"{s.first_name} {s.last_name}", req.seller_denomination) for s in req.sellers
    ] + [
        (f"{b.first_name} {b.last_name}", req.buyer_denomination) for b in req.buyers
    ]
    inject_signatures(doc, parties)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    headers = {"Content-Disposition": "attachment; filename=sell_contract.docx"}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/documents/rent_contract")
def generate_rent_contract(req: RentContractRequest):
    template_path = TEMPLATES_DIR / "rent_contract.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    # Pass 1 — docxtpl renders all text placeholders.
    # SIGNATURES is set to a sentinel so we can locate the paragraph in pass 2.
    tpl = DocxTemplate(template_path)
    context = {
        "BUFETE_NOMBRE": req.firm_name,
        "BUFETE_SERVICIO": req.firm_service,
        "NOTARIO_TITULO_NOMBRE": req.notary_title_name,
        "NOTARIO_TITULO_NOMBRE_MAYUS": req.notary_title_name.upper(),
        "NOTARIO_TELEFONO": req.notary_phone,
        "NOTARIO_MUNICIPIO": req.notary_municipality,
        "NOTARIO_MATRICULA": req.notary_registration,
        "PROPIETARIO_BLOQUE": party_block(req.owners),
        "PROPIETARIO_QUIEN": quien(len(req.owners)),
        "PROPIETARIO_SIGUE": sigue(len(req.owners)),
        "PROPIETARIO_SE_DENOMINARA": denominar(len(req.owners)),
        "DENOMINACION_PROPIETARIO": resolve_denomination(req.owner_denomination, "PROPIETARIO", req.owners),
        "INQUILINO_BLOQUE": party_block(req.tenants),
        "INQUILINO_QUIEN": quien(len(req.tenants)),
        "INQUILINO_SIGUE": sigue(len(req.tenants)),
        "INQUILINO_SE_DENOMINARA": denominar(len(req.tenants)),
        "DENOMINACION_INQUILINO": resolve_denomination(req.tenant_denomination, "INQUILINO", req.tenants),
        "DESCRIPCION_INMUEBLE": req.property_description,
        "USO_INMUEBLE": req.property_use,
        "PRECIO_LETRAS": req.price_words,
        "PRECIO_RD": req.price_rd,
        "DEPOSITO_LETRAS": req.deposit_words,
        "DEPOSITO_RD": req.deposit_rd,
        "DURACION_LETRAS": req.duration_words,
        "DURACION_NUM": req.duration_number,
        "FECHA_INICIO": req.start_date,
        "FECHA_FIN": req.end_date,
        "CANTIDAD_ORIGINALES_LETRAS": req.originals_count_words,
        "CANTIDAD_ORIGINALES": req.originals_count,
        "CIUDAD_CONTRATO": req.city,
        "PROVINCIA_CONTRATO": req.province,
        "DIA_LETRAS": req.day_words,
        "DIA_NUM": req.day_number,
        "MES": req.month,
        "ANIO_LETRAS": req.year_words,
        "ANIO_NUM": req.year_number,
        "SIGNATURES": _SIG_SENTINEL,
        "FIRMANTES_LISTA": req.signers_list,
    }
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    # Pass 2 — python-docx finds the sentinel and injects the signature table.
    doc = Document(intermediate)
    parties = [
        (f"{o.first_name} {o.last_name}", req.owner_denomination) for o in req.owners
    ] + [
        (f"{t.first_name} {t.last_name}", req.tenant_denomination) for t in req.tenants
    ]
    inject_signatures(doc, parties)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    headers = {"Content-Disposition": "attachment; filename=rent_contract.docx"}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/documents/poder_especial")
def generate_poder_especial(req: PoderEspecialRequest):
    template_path = TEMPLATES_DIR / "poder_especial.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    tpl = DocxTemplate(template_path)
    context = _notary_context(req)
    context.update(
        {
            "PODERDANTE_BLOQUE": party_block(req.poderdantes),
            "PODERDANTE_HA": ha(len(req.poderdantes)),
            "PODERDANTE_QUIEN": quien(len(req.poderdantes)),
            "PODERDANTE_SIGUE": sigue(len(req.poderdantes)),
            "PODERDANTE_SE_DENOMINARA": denominar(len(req.poderdantes)),
            "PODERDANTE_OTORGA": pick(len(req.poderdantes), "otorga", "otorgan"),
            "DENOMINACION_PODERDANTE": resolve_denomination(
                req.poderdante_denomination, "PODERDANTE", req.poderdantes
            ),
            "APODERADO_BLOQUE": party_block(req.apoderados),
            "APODERADO_QUIEN": quien(len(req.apoderados)),
            "APODERADO_SIGUE": sigue(len(req.apoderados)),
            "APODERADO_SE_DENOMINARA": denominar(len(req.apoderados)),
            "DENOMINACION_APODERADO": resolve_denomination(
                req.apoderado_denomination, "APODERADO", req.apoderados
            ),
            "OBJETO_DEL_PODER": req.power_object,
        }
    )
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    doc = Document(intermediate)
    parties = (
        [(f"{p.first_name} {p.last_name}".upper(), "Poderdante") for p in req.poderdantes]
        + [(f"{a.first_name} {a.last_name}".upper(), "Apoderado") for a in req.apoderados]
        + [(f"{w.first_name} {w.last_name}".upper(), "Testigo Instrumental") for w in req.witnesses]
        + [(req.notary_title_name.upper(), "Notario Público")]
    )
    inject_signatures(doc, parties)

    return _act_response(doc, "poder_especial.docx")


@app.post("/documents/declaracion_jurada")
def generate_declaracion_jurada(req: DeclaracionJuradaRequest):
    template_path = TEMPLATES_DIR / "declaracion_jurada.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    tpl = DocxTemplate(template_path)
    context = _notary_context(req) | _clause_act_context(req)
    context.update({"ACTO_NUMERO": req.act_number, "HORA": req.time})
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    doc = Document(intermediate)
    inject_clauses(doc, req.clauses)
    parties = (
        [(f"{d.first_name} {d.last_name}".upper(), "Compareciente") for d in req.declarants]
        + [(f"{w.first_name} {w.last_name}".upper(), "Testigo Instrumental") for w in req.witnesses]
        + [(req.notary_title_name.upper(), "Notario Público")]
    )
    inject_signatures(doc, parties)

    return _act_response(doc, "declaracion_jurada.docx")


@app.post("/documents/acto_notoriedad")
def generate_acto_notoriedad(req: ActoNotoriedadRequest):
    template_path = TEMPLATES_DIR / "acto_notoriedad.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    tpl = DocxTemplate(template_path)
    context = _notary_context(req) | _clause_act_context(req)
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    doc = Document(intermediate)
    inject_clauses(doc, req.clauses)
    parties = (
        [(f"{d.first_name} {d.last_name}".upper(), "Compareciente") for d in req.declarants]
        + [(f"{w.first_name} {w.last_name}".upper(), "Testigo Instrumental") for w in req.witnesses]
        + [(req.notary_title_name.upper(), "Notario Público")]
    )
    inject_signatures(doc, parties)

    return _act_response(doc, "acto_notoriedad.docx")


@app.post("/documents/determinacion_herederos")
def generate_determinacion_herederos(req: DeterminacionHerederosRequest):
    template_path = TEMPLATES_DIR / "determinacion_herederos.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    tpl = DocxTemplate(template_path)
    context = _notary_context(req)
    context.update(
        {
            "COMPARECIENTE_BLOQUE": party_block(req.declarants),
            "COMPARECIENTE_HA": ha(len(req.declarants)),
            "COMPARECIENTE_QUIEN": quien(len(req.declarants)),
            "COMPARECIENTE_SIGUE": sigue(len(req.declarants)),
            "COMPARECIENTE_SE_DENOMINARA": denominar(len(req.declarants)),
            "DENOMINACION_COMPARECIENTE": resolve_denomination(
                req.declarant_denomination, "DECLARANTE", req.declarants
            ),
            "DIFUNTO_NOMBRE": _bold_upper(req.deceased_name),
            "DIFUNTO_NACIONALIDAD": req.deceased_nationality,
            "DIFUNTO_ESTADO_CIVIL": req.deceased_marital_status,
            "DIFUNTO_OCUPACION": req.deceased_occupation,
            "DIFUNTO_CEDULA": req.deceased_id_number,
            "DIFUNTO_FECHA_NACIMIENTO": req.deceased_birth_date,
            "DIFUNTO_FECHA_DEFUNCION": req.deceased_death_date,
            "DIFUNTO_ACTA_DEFUNCION": req.deceased_death_record,
            "DIFUNTO_LUGAR_ORIGEN": req.deceased_origin_place,
            "DIFUNTO_CONYUGE": req.deceased_spouse_name,
            "HEREDEROS": _HEIRS_SENTINEL,
            "CLAUSULAS": _CLAUSES_SENTINEL,
        }
    )
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    doc = Document(intermediate)
    inject_heirs(doc, req.heirs)
    # PRIMERO (deceased facts) and SEGUNDO (heirs enumeration) are fixed in the
    # template; dynamic clauses continue numbering from TERCERO (index 2).
    inject_clauses(doc, req.clauses, start_index=2)
    parties = (
        [(f"{d.first_name} {d.last_name}".upper(), "Compareciente") for d in req.declarants]
        + [(f"{w.first_name} {w.last_name}".upper(), "Testigo Instrumental") for w in req.witnesses]
        + [(req.notary_title_name.upper(), "Notario Público")]
    )
    inject_signatures(doc, parties)

    return _act_response(doc, "determinacion_herederos.docx")


@app.post("/documents/autorizacion_viaje_menor")
def generate_autorizacion_viaje_menor(req: AutorizacionViajeMenorRequest):
    template_path = TEMPLATES_DIR / "autorizacion_viaje_menor.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    tpl = DocxTemplate(template_path)
    context = _notary_context(req)
    context.update(
        {
            "AUTORIZADOR_BLOQUE": party_block(req.authorizers),
            "AUTORIZADOR_HA": ha(len(req.authorizers)),
            "AUTORIZADOR_QUIEN": quien(len(req.authorizers)),
            "AUTORIZADOR_SIGUE": sigue(len(req.authorizers)),
            "AUTORIZADOR_SE_DENOMINARA": denominar(len(req.authorizers)),
            "AUTORIZADOR_DA": pick(len(req.authorizers), "da", "dan"),
            "DENOMINACION_AUTORIZADOR": resolve_denomination(
                req.authorizer_denomination, "AUTORIZANTE", req.authorizers
            ),
            "MENOR_NOMBRE": _bold_upper(req.minor_name),
            "MENOR_FECHA_NACIMIENTO": req.minor_birth_date,
            "MENOR_ACTA_NACIMIENTO": req.minor_birth_record,
            "MENOR_PASAPORTE": req.minor_passport,
            "PAIS_DESTINO": req.destination_country,
            "ACEPTANTE_BLOQUE": party_block(req.acceptors),
            "ACEPTANTE_ACEPTA": pick(len(req.acceptors), "acepta", "aceptan"),
            "ACEPTANTE_QUIEN": quien(len(req.acceptors)),
            "ACEPTANTE_SIGUE": sigue(len(req.acceptors)),
            "ACEPTANTE_SE_DENOMINARA": denominar(len(req.acceptors)),
            "DENOMINACION_ACEPTANTE": resolve_denomination(
                req.acceptor_denomination, "ACEPTANTE", req.acceptors
            ),
        }
    )
    tpl.render(context)
    intermediate = BytesIO()
    tpl.save(intermediate)
    intermediate.seek(0)

    doc = Document(intermediate)
    parties = (
        [(f"{a.first_name} {a.last_name}".upper(), "Autorizador") for a in req.authorizers]
        + [(f"{a.first_name} {a.last_name}".upper(), "Aceptante") for a in req.acceptors]
        + [(f"{w.first_name} {w.last_name}".upper(), "Testigo Instrumental") for w in req.witnesses]
        + [(req.notary_title_name.upper(), "Notario Público")]
    )
    inject_signatures(doc, parties)

    return _act_response(doc, "autorizacion_viaje_menor.docx")


# Static files last so API routes take precedence
if STATIC_DIR.exists():
    app.mount("/ui", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")
